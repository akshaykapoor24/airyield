"""
DealExtractionService
─────────────────────
Parses uploaded PDF / Excel / Word / image deal files from travel agents
and returns a list of structured DealRow dicts for user review.

Supported formats
  • Excel (.xls / .xlsx)  — openpyxl + pandas
  • PDF                   — PyMuPDF (fitz) table + text extraction
  • Word (.docx)          — python-docx  (optional; falls back gracefully)
  • Image                 — Pillow (OCR placeholder)

Returned structure
  {
      "source_type": "excel" | "pdf" | "word" | "image",
      "file_name": str,
      "rows": [
          {
              "row_order": int,
              "airline_name": str,
              "iata_code": str,
              "variant": str,
              "eco_commission": str,
              "peco_commission": str,
              "bus_commission": str,
              "base_type": str,
              "valid_on": str,
              "valid_from": str | None,   # ISO date string
              "valid_to": str | None,
              "validity_raw": str,
              "remarks": str,
          },
          ...
      ],
      "confidence": float,   # 0.0 – 1.0
  }
"""
from __future__ import annotations

import io
import re
import logging
from datetime import date, datetime
from typing import Any

from fastapi import UploadFile

logger = logging.getLogger(__name__)

# ── cell cleaner ───────────────────────────────────────────────────────────

def _clean_iata(raw: str | None) -> str:
    """
    Extract only the real IATA airline code (2–3 uppercase letters/digits)
    from a potentially multi-line cell like 'AI\n(LORDS ISSUANCE)'.
    """
    if not raw:
        return ""
    # take first line only
    first_line = str(raw).split("\n")[0].strip()
    # extract the first 2-3 char word that looks like an IATA code
    m = re.match(r"^([A-Z0-9]{2,3})\b", first_line)
    return m.group(1) if m else first_line[:10].strip()


def _clean_airline(raw: str | None) -> str:
    """
    Clean airline name: take only the first meaningful line,
    remove parenthetical validity info like '(OB/IB till 31 Mar 2026)'.
    """
    if not raw:
        return ""
    # take first line
    first_line = str(raw).split("\n")[0].strip()
    # remove trailing validity / date patterns like (OB/IB till 31 Mar 2026)
    cleaned = re.sub(r"\s*\(\s*(OB|IB|OB/IB)[^)]*\)", "", first_line, flags=re.IGNORECASE).strip()
    # remove leading/trailing punctuation
    cleaned = cleaned.strip(".,;:-").strip()
    return cleaned or first_line


def _clean_str(raw: str | None, max_len: int = 255) -> str:
    """Generic string cleaner — collapse newlines and trim to max_len."""
    if not raw:
        return ""
    cleaned = " | ".join(line.strip() for line in str(raw).splitlines() if line.strip())
    return cleaned[:max_len]


def _clean_row(row: dict) -> dict:
    """Apply all field-level cleaners to an extracted row dict."""
    row["airline_name"]   = _clean_airline(row.get("airline_name"))
    row["iata_code"]      = _clean_iata(row.get("iata_code"))
    row["variant"]        = _clean_str(row.get("variant"), 100)
    row["eco_commission"] = _clean_str(row.get("eco_commission"), 50)
    row["peco_commission"]= _clean_str(row.get("peco_commission"), 50)
    row["bus_commission"] = _clean_str(row.get("bus_commission"), 50)
    row["base_type"]      = _clean_str(row.get("base_type"), 20)
    row["valid_on"]       = _clean_str(row.get("valid_on"), 20)
    row["validity_raw"]   = _clean_str(row.get("validity_raw"), 255)
    return row


# ── helpers ────────────────────────────────────────────────────────────────


# Column alias → standard key
_COL_MAP: dict[str, str] = {
    # airline
    "airline": "airline_name",
    "airline name": "airline_name",
    "carrier": "airline_name",
    # iata
    "airline code": "iata_code",
    "code": "iata_code",
    "iata": "iata_code",
    "tkt. iata": "iata_code",
    "tkt.iata": "iata_code",
    # economy
    "eco": "eco",
    "economy": "eco",
    "p.eco": "peco",
    "p eco": "peco",
    "p.ecom": "peco",
    "pecom": "peco",
    "peco": "peco",
    "pey": "peco",
    "premium economy": "peco",
    # business
    "bus": "bus",
    "business": "bus",
    "buss": "bus",
    "buss/first": "bus",
    "business/first": "bus",
    "first": "bus",
    # validity
    "validity": "validity_raw",
    "valid on": "valid_on",
    "valid till": "validity_raw",
    "validon": "valid_on",
    # remarks
    "remarks": "remarks",
    "plb remarks & conditions": "remarks",
    "plb remarks": "remarks",
    "conditions": "remarks",
    "remark": "remarks",
    # base type
    "base or basse + f.s.": "base_type",
    "base": "base_type",
    "base or base + f.s.": "base_type",
}

_MONTH_MAP = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}


def _normalise_col(col: str) -> str:
    return _COL_MAP.get(col.strip().lower(), col.strip().lower())


def _parse_date(text: str | None) -> date | None:
    if not text:
        return None
    text = str(text).strip()
    # patterns: "31 Mar 2026", "31-Mar-2026", "2026-03-31", "31/03/2026"
    patterns = [
        r"(\d{1,2})[\s\-/]([A-Za-z]{3})[\s\-/](\d{4})",   # 31 Mar 2026
        r"(\d{4})-(\d{2})-(\d{2})",                          # 2026-03-31
        r"(\d{1,2})/(\d{1,2})/(\d{4})",                      # 31/03/2026
        r"(\d{1,2})-(\d{1,2})-(\d{4})",                      # 31-03-2026
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            try:
                g = m.groups()
                if len(g[1]) == 3 and g[1].lower() in _MONTH_MAP:   # "31 Mar 2026"
                    return date(int(g[2]), _MONTH_MAP[g[1].lower()], int(g[0]))
                elif len(g[0]) == 4:                                   # "2026-03-31"
                    return date(int(g[0]), int(g[1]), int(g[2]))
                else:
                    return date(int(g[2]), int(g[1]), int(g[0]))
            except Exception:
                continue
    return None


def _parse_validity(raw: str) -> tuple[str | None, str | None, str | None]:
    """
    Parse a validity string like "OB/IB till 31 Mar 2026" or "OB 1Apr 26 / IB till 31 Mar 2027".
    Returns (valid_on, valid_from_iso, valid_to_iso).
    """
    if not raw:
        return None, None, None
    raw_str = str(raw).strip()

    valid_on = None
    if re.search(r"\bOB/IB\b", raw_str, re.I):
        valid_on = "OB/IB"
    elif re.search(r"\bOB\b", raw_str, re.I):
        valid_on = "OB"
    elif re.search(r"\bIB\b", raw_str, re.I):
        valid_on = "IB"
    elif re.search(r"\bB\+YQ\b", raw_str, re.I):
        valid_on = "B+YQ"

    # look for two dates (from / to)
    dates_found = re.findall(
        r"\b(\d{1,2}[\s\-/][A-Za-z]{3}[\s\-/]\d{2,4}|\d{4}-\d{2}-\d{2}|\d{1,2}[A-Za-z]{3}\s*\d{2,4})\b",
        raw_str, re.I
    )
    parsed_dates = [_parse_date(d) for d in dates_found if _parse_date(d)]
    parsed_dates = [d for d in parsed_dates if d is not None]

    valid_from = str(parsed_dates[0]) if len(parsed_dates) >= 1 else None
    valid_to   = str(parsed_dates[-1]) if len(parsed_dates) >= 1 else None

    # "till 31 Mar 2026" — only one date → that is the end date
    if len(parsed_dates) == 1 and re.search(r"\btill\b", raw_str, re.I):
        valid_from = None
        valid_to   = str(parsed_dates[0])

    return valid_on, valid_from, valid_to


def _parse_commission(raw: str | None) -> tuple[str, str]:
    """
    Returns (commission_str, base_type).
    e.g. "2.75% (B+YQ)"  → ("2.75%", "B+YQ")
         "2.75"           → ("2.75%", "")
         0.0175 (float)   → ("1.75%", "")   ← Excel stores % as decimal
         0.0 / "0" / "-"  → ("", "")
    """
    if raw is None:
        return "", ""

    # numeric path: Excel stores 1.75% as float 0.0175
    if isinstance(raw, (int, float)):
        val = float(raw)
        if val == 0.0:
            return "", ""
        # if stored as decimal fraction (≤ 1.0), convert to percentage
        pct = val * 100 if val <= 1.0 else val
        return f"{pct:.2f}%".rstrip("0").rstrip(".") + "%", ""

    raw_str = str(raw).strip()
    if raw_str in ("-", "", "nan", "None", "0", "0.0", "0%"):
        return "", ""

    base = ""
    for bt in ("B+YQ", "B+YR", "B+YQ+YR", "RTT/RH", "B"):
        if bt in raw_str.upper():
            base = bt
            break
    # explicit percentage in string
    m = re.search(r"([\d.]+)\s*%", raw_str)
    if m:
        commission = m.group(1) + "%"
    else:
        m2 = re.search(r"^([\d.]+)$", raw_str)
        if m2:
            val = float(m2.group(1))
            # decimal fraction → convert
            pct = val * 100 if val > 0 and val <= 1.0 else val
            commission = f"{pct:.4f}".rstrip("0").rstrip(".") + "%"
        else:
            commission = raw_str
    return commission, base


# ── Excel parser ───────────────────────────────────────────────────────────

_COMMISSION_KEYWORDS = {"eco", "economy", "pey", "p.eco", "p eco", "peco",
                         "bus", "buss", "business", "first", "buss/first"}
_IDENTITY_KEYWORDS   = {"airline", "carrier", "code", "validity", "remarks",
                         "iata", "base", "valid"}


def _row_has_keywords(row_str: list[str], kw_set: set[str], min_count: int = 2) -> bool:
    return sum(1 for cell in row_str if any(kw in cell for kw in kw_set)) >= min_count


def _build_merged_header(row_a: list[str], row_b: list[str]) -> list[str]:
    """
    Merge a two-level header where row_a has identity cols + period labels
    and row_b has the commission sub-column names (ECO, PEY, Buss).
    Result: take row_b where it has a commission keyword, else row_a.
    """
    merged = []
    for a, b in zip(row_a, row_b):
        b_lower = b.strip().lower()
        if any(kw in b_lower for kw in _COMMISSION_KEYWORDS) and b_lower:
            merged.append(b)
        elif a.strip():
            merged.append(a)
        else:
            merged.append(b)
    return merged


def _extract_excel(content: bytes, filename: str) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    ws = wb.active

    all_rows = list(ws.iter_rows(values_only=True))
    rows: list[dict] = []

    # ── find header row(s) ────────────────────────────────────────────────
    # Strategy:
    #  1. Find the BEST header row — one that has BOTH identity cols AND commission cols
    #  2. If a row has identity cols but not commission cols, look at the NEXT row for
    #     commission sub-headers (two-level header like Lords/PLB sheets)
    header_idx = -1
    header_row: list[str] = []
    orig_header_row: list[str] = []  # original (un-lowercased) for doc_columns

    all_keywords = _IDENTITY_KEYWORDS | _COMMISSION_KEYWORDS

    for i, row in enumerate(all_rows):
        row_str  = [str(c).strip().lower() if c is not None else "" for c in row]
        orig_str = [str(c).strip()         if c is not None else "" for c in row]

        has_identity   = _row_has_keywords(row_str, _IDENTITY_KEYWORDS, 2)
        has_commission = _row_has_keywords(row_str, _COMMISSION_KEYWORDS, 1)

        if has_identity and has_commission:
            # perfect single-row header
            header_idx = i
            header_row = row_str
            orig_header_row = orig_str
            break

        if has_identity and not has_commission and i + 1 < len(all_rows):
            # possible two-level header — check the next row for commission sub-cols
            next_str  = [str(c).strip().lower() if c is not None else "" for c in all_rows[i + 1]]
            orig_next = [str(c).strip()         if c is not None else "" for c in all_rows[i + 1]]
            if _row_has_keywords(next_str, _COMMISSION_KEYWORDS, 1):
                # merge the two rows
                header_idx = i + 1   # data starts after the sub-header row
                header_row = _build_merged_header(row_str, next_str)
                orig_header_row = _build_merged_header(orig_str, orig_next)
                break

    if header_idx == -1:
        return _extract_excel_pandas(content, filename)

    # ── build doc_columns + raw_rows from original headers ────────────────
    seen_orig: dict[str, int] = {}
    unique_orig: list[str] = []
    for h in orig_header_row:
        if not h:
            unique_orig.append("")
            continue
        if h in seen_orig:
            seen_orig[h] += 1
            unique_orig.append(f"{h} ({seen_orig[h]})")
        else:
            seen_orig[h] = 1
            unique_orig.append(h)

    doc_columns: list[str] = [h for h in unique_orig if h]

    raw_rows_data: list[dict] = []
    for orig_row in all_rows[header_idx + 1:]:
        if all(c is None or str(c).strip() == "" for c in orig_row):
            continue
        rd: dict[str, str] = {}
        for col, val in zip(unique_orig, orig_row):
            if col:
                rd[col] = str(val).strip() if val is not None else ""
        if any(v for v in rd.values()):
            raw_rows_data.append(rd)

    # ── normalise headers ─────────────────────────────────────────────────
    # handle multi-period: detect duplicate eco/bus columns
    std_headers: list[str] = []
    eco_count = bus_count = peco_count = 0
    for h in header_row:
        std = _normalise_col(h)
        if std == "eco":
            eco_count += 1
            std_headers.append(f"eco_{eco_count}")
        elif std == "peco":
            peco_count += 1
            std_headers.append(f"peco_{peco_count}")
        elif std == "bus":
            bus_count += 1
            std_headers.append(f"bus_{bus_count}")
        else:
            std_headers.append(std)

    n_periods = max(eco_count, bus_count, 1)

    # ── parse data rows ───────────────────────────────────────────────────
    last_airline = ""
    last_iata    = ""
    order = 0

    for raw_row in all_rows[header_idx + 1:]:
        if all(c is None or str(c).strip() == "" for c in raw_row):
            continue  # skip empty rows

        # keep raw types (float, int) so _parse_commission can handle decimal percentages
        row_dict: dict[str, Any] = dict(zip(std_headers, [c if c is not None else "" for c in raw_row]))

        # carry forward airline / iata from merged cells
        airline = str(row_dict.get("airline_name", "") or "").strip()
        iata    = str(row_dict.get("iata_code", "")    or "").strip()
        if airline and airline.lower() not in ("nan", "none"):
            last_airline = airline
        else:
            airline = last_airline
        if iata and iata.lower() not in ("nan", "none"):
            last_iata = iata
        else:
            iata = last_iata

        # skip rows that look like sub-headers
        if airline.lower() in ("airline", "carrier", "airline name"):
            continue

        validity_raw = str(row_dict.get("validity_raw", "") or "")
        valid_on_raw = str(row_dict.get("valid_on", "") or "")
        valid_on_parsed, vf, vt = _parse_validity(validity_raw or valid_on_raw)

        base_type_raw = str(row_dict.get("base_type", "") or "")

        for period in range(1, n_periods + 1):
            eco_raw  = row_dict.get(f"eco_{period}",  row_dict.get("eco_1", row_dict.get("eco", "")))
            peco_raw = row_dict.get(f"peco_{period}", row_dict.get("peco_1", row_dict.get("peco", "")))
            bus_raw  = row_dict.get(f"bus_{period}",  row_dict.get("bus_1", row_dict.get("bus", "")))

            eco_com, eco_base  = _parse_commission(eco_raw)
            peco_com, peco_base = _parse_commission(peco_raw)
            bus_com, bus_base   = _parse_commission(bus_raw)

            # skip rows with no commission data
            if not any([eco_com, peco_com, bus_com]):
                continue

            base = eco_base or peco_base or bus_base or base_type_raw or ""

            rows.append({
                "row_order":       order,
                "airline_name":    airline,
                "iata_code":       iata,
                "variant":         _guess_variant(airline),
                "eco_commission":  eco_com,
                "peco_commission": peco_com,
                "bus_commission":  bus_com,
                "base_type":       base,
                "valid_on":        valid_on_parsed or valid_on_raw or "",
                "valid_from":      vf,
                "valid_to":        vt,
                "validity_raw":    validity_raw or valid_on_raw or "",
                "remarks":         str(row_dict.get("remarks", "") or ""),
            })
            order += 1
            if n_periods == 1:
                break   # no multi-period splitting needed

    rows = [_clean_row(r) for r in rows]
    confidence = 0.9 if rows else 0.3
    return {
        "source_type": "excel", "file_name": filename,
        "rows": rows, "confidence": confidence,
        "doc_columns": doc_columns, "raw_rows": raw_rows_data,
    }


def _extract_excel_pandas(content: bytes, filename: str) -> dict:
    """Fallback Excel extraction via pandas when openpyxl header detection fails."""
    import pandas as pd
    df = pd.read_excel(io.BytesIO(content), header=None)
    rows = []
    order = 0
    for _, row in df.iterrows():
        cells = [str(c).strip() for c in row if str(c).strip() not in ("nan", "")]
        if len(cells) >= 3:
            rows.append({
                "row_order":       order,
                "airline_name":    cells[0] if len(cells) > 0 else "",
                "iata_code":       cells[1] if len(cells) > 1 else "",
                "variant":         "",
                "eco_commission":  cells[2] if len(cells) > 2 else "",
                "peco_commission": cells[3] if len(cells) > 3 else "",
                "bus_commission":  cells[4] if len(cells) > 4 else "",
                "base_type":       "",
                "valid_on":        "",
                "valid_from":      None,
                "valid_to":        None,
                "validity_raw":    cells[5] if len(cells) > 5 else "",
                "remarks":         " | ".join(cells[6:]) if len(cells) > 6 else "",
            })
            order += 1
    return {"source_type": "excel", "file_name": filename, "rows": rows, "confidence": 0.5,
            "doc_columns": [], "raw_rows": []}


def _guess_variant(airline_name: str) -> str:
    """Derive variant from airline name string e.g. 'AIR INDIA (INTL)' → 'INTL'."""
    m = re.search(r"\(([^)]+)\)", airline_name)
    return m.group(1).strip() if m else ""


# ── PDF parser ─────────────────────────────────────────────────────────────

def _extract_pdf(content: bytes, filename: str) -> dict:
    import fitz  # PyMuPDF

    doc  = fitz.open(stream=content, filetype="pdf")
    rows: list[dict] = []
    order = 0
    doc_columns: list[str] = []
    raw_rows_data: list[dict] = []

    for page in doc:
        # try structured table extraction first
        tabs = page.find_tables()
        if tabs and tabs.tables:
            for tab in tabs.tables:
                table_data = tab.extract()
                if not table_data:
                    continue

                # find header row within this table
                header_idx = -1
                headers: list[str] = []
                orig_headers: list[str] = []
                for i, trow in enumerate(table_data):
                    row_str = [str(c).strip().lower() if c else "" for c in trow]
                    keywords = {"airline", "eco", "bus", "validity", "code", "commission"}
                    if sum(1 for cell in row_str if any(kw in cell for kw in keywords)) >= 2:
                        header_idx = i
                        headers      = [_normalise_col(c) for c in row_str]
                        orig_headers = [str(c).strip() if c else "" for c in trow]
                        break

                if header_idx == -1:
                    continue

                # Capture doc_columns + raw_rows from the first valid table
                if not doc_columns and orig_headers:
                    doc_columns = [h for h in orig_headers if h]
                    for trow in table_data[header_idx + 1:]:
                        if not trow or all(not c for c in trow):
                            continue
                        rd: dict[str, str] = {}
                        for col, val in zip(orig_headers, trow):
                            if col:
                                rd[col] = str(val).strip() if val else ""
                        if any(v for v in rd.values()):
                            raw_rows_data.append(rd)

                for trow in table_data[header_idx + 1:]:
                    if not trow or all(not c for c in trow):
                        continue
                    row_dict = dict(zip(headers, [str(c).strip() if c else "" for c in trow]))

                    airline = row_dict.get("airline_name", "")
                    iata    = row_dict.get("iata_code", "")
                    val_raw = row_dict.get("validity_raw", "")
                    valid_on_raw = row_dict.get("valid_on", "")
                    valid_on, vf, vt = _parse_validity(val_raw or valid_on_raw)

                    eco_raw  = row_dict.get("eco", "")
                    peco_raw = row_dict.get("peco", "")
                    bus_raw  = row_dict.get("bus", "")

                    eco_com, eco_base   = _parse_commission(eco_raw)
                    peco_com, _         = _parse_commission(peco_raw)
                    bus_com, bus_base   = _parse_commission(bus_raw)

                    if not any([eco_com, peco_com, bus_com, airline]):
                        continue

                    rows.append({
                        "row_order":       order,
                        "airline_name":    airline,
                        "iata_code":       iata,
                        "variant":         _guess_variant(airline),
                        "eco_commission":  eco_com,
                        "peco_commission": peco_com,
                        "bus_commission":  bus_com,
                        "base_type":       eco_base or bus_base or "",
                        "valid_on":        valid_on or valid_on_raw or "",
                        "valid_from":      vf,
                        "valid_to":        vt,
                        "validity_raw":    val_raw or valid_on_raw or "",
                        "remarks":         row_dict.get("remarks", ""),
                    })
                    order += 1
        else:
            # fallback: plain text line-by-line
            text = page.get_text()
            extracted = _parse_text_lines(text, order)
            rows.extend(extracted)
            order += len(extracted)

    rows = [_clean_row(r) for r in rows]
    confidence = 0.75 if rows else 0.2
    return {
        "source_type": "pdf", "file_name": filename,
        "rows": rows, "confidence": confidence,
        "doc_columns": doc_columns, "raw_rows": raw_rows_data,
    }


def _parse_text_lines(text: str, start_order: int) -> list[dict]:
    """
    Heuristic: scan text lines for percentage patterns like "2.00% (B+YQ)".
    Groups them by the nearest airline name line above.
    """
    rows = []
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    order = start_order
    current_airline = ""
    current_iata    = ""

    for line in lines:
        # airline line detection: all caps, no %, contains letters
        if re.match(r"^[A-Z][A-Z\s()./&\-]+$", line) and "%" not in line and len(line) > 3:
            # could be an airline name
            iata_m = re.match(r"^([A-Z]{2})\s+(.+)$", line)
            if iata_m:
                current_iata    = iata_m.group(1)
                current_airline = iata_m.group(2).strip()
            else:
                current_airline = line
                current_iata    = ""
            continue

        # commission line: contains at least one percentage
        commissions = re.findall(r"([\d.]+\s*%\s*(?:\([^)]+\))?)", line)
        if commissions and current_airline:
            eco_com, eco_base   = _parse_commission(commissions[0] if len(commissions) > 0 else "")
            peco_com, _         = _parse_commission(commissions[1] if len(commissions) > 1 else "")
            bus_com, bus_base   = _parse_commission(commissions[2] if len(commissions) > 2 else "")

            # validity from same line
            valid_on, vf, vt = _parse_validity(line)

            rows.append({
                "row_order":       order,
                "airline_name":    current_airline,
                "iata_code":       current_iata,
                "variant":         _guess_variant(current_airline),
                "eco_commission":  eco_com,
                "peco_commission": peco_com,
                "bus_commission":  bus_com,
                "base_type":       eco_base or bus_base or "",
                "valid_on":        valid_on or "",
                "valid_from":      vf,
                "valid_to":        vt,
                "validity_raw":    line,
                "remarks":         "",
            })
            order += 1

    return rows


# ── Word parser ─────────────────────────────────────────────────────────────

def _extract_word(content: bytes, filename: str) -> dict:
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        rows = []
        order = 0
        doc_columns: list[str] = []
        raw_rows_data: list[dict] = []
        for table in doc.tables:
            header_row_idx = -1
            headers: list[str] = []
            orig_headers: list[str] = []
            for i, row in enumerate(table.rows):
                cells = [c.text.strip().lower() for c in row.cells]
                keywords = {"airline", "eco", "bus", "validity", "code"}
                if sum(1 for c in cells if any(kw in c for kw in keywords)) >= 2:
                    header_row_idx = i
                    headers = [_normalise_col(c) for c in cells]
                    orig_headers = [c.text.strip() for c in row.cells]
                    break

            if header_row_idx == -1:
                continue

            # Capture doc_columns + raw_rows from first valid table
            if not doc_columns and orig_headers:
                doc_columns = [h for h in orig_headers if h]
                for row in table.rows[header_row_idx + 1:]:
                    rd = {orig_headers[j]: c.text.strip() for j, c in enumerate(row.cells) if j < len(orig_headers) and orig_headers[j]}
                    if any(v for v in rd.values()):
                        raw_rows_data.append(rd)

            for row in table.rows[header_row_idx + 1:]:
                cells = [c.text.strip() for c in row.cells]
                row_dict = dict(zip(headers, cells))
                airline = row_dict.get("airline_name", "")
                val_raw = row_dict.get("validity_raw", "")
                valid_on, vf, vt = _parse_validity(val_raw)
                eco_com, eco_base  = _parse_commission(row_dict.get("eco", ""))
                peco_com, _        = _parse_commission(row_dict.get("peco", ""))
                bus_com, bus_base  = _parse_commission(row_dict.get("bus", ""))
                if not any([eco_com, peco_com, bus_com, airline]):
                    continue
                rows.append({
                    "row_order":       order,
                    "airline_name":    airline,
                    "iata_code":       row_dict.get("iata_code", ""),
                    "variant":         _guess_variant(airline),
                    "eco_commission":  eco_com,
                    "peco_commission": peco_com,
                    "bus_commission":  bus_com,
                    "base_type":       eco_base or bus_base or "",
                    "valid_on":        valid_on or "",
                    "valid_from":      vf,
                    "valid_to":        vt,
                    "validity_raw":    val_raw,
                    "remarks":         row_dict.get("remarks", ""),
                })
                order += 1
        rows = [_clean_row(r) for r in rows]
        return {
            "source_type": "word", "file_name": filename,
            "rows": rows, "confidence": 0.75,
            "doc_columns": doc_columns, "raw_rows": raw_rows_data,
        }
    except Exception as e:
        logger.warning(f"Word extraction failed: {e}")
        return {"source_type": "word", "file_name": filename, "rows": [], "confidence": 0.0,
                "doc_columns": [], "raw_rows": []}


# ── Image parser ────────────────────────────────────────────────────────────

def _extract_image(content: bytes, filename: str) -> dict:
    # pytesseract optional — return empty rows with note if not installed
    try:
        import pytesseract
        from PIL import Image
        img  = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(img)
        rows = _parse_text_lines(text, 0)
        return {"source_type": "image", "file_name": filename, "rows": rows, "confidence": 0.5,
                "doc_columns": [], "raw_rows": []}
    except ImportError:
        return {
            "source_type": "image",
            "file_name": filename,
            "rows": [],
            "confidence": 0.0,
            "warning": "OCR not available. Install pytesseract to extract image deals.",
            "doc_columns": [], "raw_rows": [],
        }


# ── Public API ──────────────────────────────────────────────────────────────

class DealExtractionService:

    @staticmethod
    async def extract(file: UploadFile) -> dict:
        content  = await file.read()
        filename = file.filename or "upload"
        ext      = filename.rsplit(".", 1)[-1].lower()

        if ext in ("xls", "xlsx"):
            return _extract_excel(content, filename)
        elif ext == "pdf":
            return _extract_pdf(content, filename)
        elif ext in ("doc", "docx"):
            return _extract_word(content, filename)
        elif ext in ("png", "jpg", "jpeg", "bmp", "tiff"):
            return _extract_image(content, filename)
        else:
            return {
                "source_type": "unknown",
                "file_name": filename,
                "rows": [],
                "confidence": 0.0,
                "warning": f"Unsupported file type: .{ext}",
            }
